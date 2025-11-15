from docx import Document
from pathlib import Path
from typing import List
from ..models.data_models import TextBlock, TextBlockType, ParsedDocument
import re


class DOCXDocumentParser:
    def parse_docx(self, file_path: str) -> ParsedDocument:
        """Парсит DOCX файл и возвращает структурированный документ"""
        print(f"Парсим DOCX: {file_path}")

        doc = Document(file_path)
        blocks = []
        current_page = 1
        line_count = 0
        LINES_PER_PAGE = 50  # Примерное количество строк на страницу

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Увеличиваем номер страницы каждые N строк
                line_count += 1
                if line_count >= LINES_PER_PAGE:
                    current_page += 1
                    line_count = 0

                block_type = self._classify_paragraph(paragraph, text)

                text_block = TextBlock(
                    text=text,
                    block_type=block_type,
                    page_num=current_page,
                    bbox=(0, 0, 100, 100),
                    font_size=self._get_font_size(paragraph),
                    is_bold=self._is_bold(paragraph),
                    is_italic=self._is_italic(paragraph)
                )
                blocks.append(text_block)

        # Обрабатываем таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        line_count += 1
                        if line_count >= LINES_PER_PAGE:
                            current_page += 1
                            line_count = 0

                        block_type = self._classify_text(text)
                        text_block = TextBlock(
                            text=text,
                            block_type=block_type,
                            page_num=current_page,
                            bbox=(0, 0, 100, 100)
                        )
                        blocks.append(text_block)

        print(f"DOCX распарсен: {len(blocks)} блоков, {current_page} страниц")
        return ParsedDocument(
            metadata={"format": "DOCX", "file": Path(file_path).name, "pages": current_page},
            main_content=blocks,
            raw_text="\n".join(block.text for block in blocks)
        )

    def _classify_paragraph(self, paragraph, text: str) -> TextBlockType:
        """Классифицирует параграф DOCX по типу"""
        text_lower = text.lower()

        # Заголовки по стилю
        if paragraph.style.name.startswith('Heading'):
            return TextBlockType.HEADING

        # Заголовки по содержанию
        if any(keyword in text_lower for keyword in [
            'список используемых источников', 'библиография', 'литература',
            'references', 'bibliography', 'аннотация', 'введение',
            'заключение', 'abstract', 'introduction', 'conclusion'
        ]):
            return TextBlockType.HEADING

        # Библиографические записи
        if self._is_bibliography_entry(text):
            return TextBlockType.BIBLIOGRAPHY

        return TextBlockType.PARAGRAPH

    def _classify_text(self, text: str) -> TextBlockType:
        """Классифицирует произвольный текст"""
        text_lower = text.lower()

        if any(keyword in text_lower for keyword in [
            'список используемых источников', 'библиография', 'references'
        ]):
            return TextBlockType.HEADING

        if self._is_bibliography_entry(text):
            return TextBlockType.BIBLIOGRAPHY

        return TextBlockType.PARAGRAPH

    def _is_bibliography_entry(self, text: str) -> bool:
        """Проверяет, похож ли текст на библиографическую запись"""
        # Начинается с цифры и точки
        if re.match(r'^\d+\.\s+', text):
            return True

        # Содержит типичные библиографические элементы
        if any(keyword in text.lower() for keyword in [
            'изд-во', 'издательство', 'университет', 'год', 'с.', 'сс.'
        ]) and any(str(year) in text for year in range(2000, 2025)):
            return True

        return False

    def _get_font_size(self, paragraph) -> float:
        """Получает размер шрифта параграфа"""
        try:
            if paragraph.runs:
                return paragraph.runs[0].font.size.pt if paragraph.runs[0].font.size else 12
        except:
            pass
        return 12.0

    def _is_bold(self, paragraph) -> bool:
        """Проверяет жирный шрифт"""
        try:
            if paragraph.runs:
                return paragraph.runs[0].bold or False
        except:
            pass
        return False

    def _is_italic(self, paragraph) -> bool:
        """Проверяет курсивный шрифт"""
        try:
            if paragraph.runs:
                return paragraph.runs[0].italic or False
        except:
            pass
        return False